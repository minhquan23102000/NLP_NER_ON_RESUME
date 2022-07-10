from io import BytesIO, StringIO

import docx
import fitz
from preprocessor import preprocessor
from streamlit.uploaded_file_manager import UploadedFile


def read_pdf(file):
    from pdfminer.converter import TextConverter
    from pdfminer.layout import LAParams
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfinterp import PDFPageInterpreter, PDFResourceManager
    from pdfminer.pdfpage import PDFPage
    from pdfminer.pdfparser import PDFParser

    output_string = StringIO()
    parser = PDFParser(file)
    doc = PDFDocument(parser)
    rsrcmgr = PDFResourceManager()
    device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    for page in PDFPage.create_pages(doc):
        interpreter.process_page(page)

    return output_string.getvalue()

class PDFReader(object):
    def __init__(self, need_clean=True):
        self.need_clean = need_clean
        pass

    def extract_text_from_pdf(self, file: UploadedFile):
        """Extract raw text content from pdf file

        Args:
            filename (_type_): _description_
        """
        pdf = fitz.open(stream=file.read(), filetype="pdf")
        cv_text = ""
        for page in pdf:
            cv_text += page.get_text() + '\n'

        return cv_text

    def extract_text_from_doc(self, file):
        document = docx.Document(file)
        text = ""
        for p in document.paragraphs:
            text += "\n" + p.text

        return text

    def read(self, resume_file: UploadedFile, fast=True):
        text = ""
        mime_type = resume_file.name.split(".")[-1]
        try:
            if mime_type == "pdf":
                if fast:
                    text = self.extract_text_from_pdf(resume_file)
                else:
                    text = read_pdf(resume_file)
            elif mime_type == "docx":
                text = self.extract_text_from_doc(resume_file)
        except Exception as e:
            print(f"Can't read file {resume_file.name}, cause: {str(e)}")
            return ""

        if self.need_clean:
            text = preprocessor.clean_text(text)

        return text
